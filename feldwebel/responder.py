"""
DeepSeek response engine with native tool use (function calling).
The Feldwebel — operational assistant for the Kommandant.

Architecture: DeepSeek decides when to call tools (gmail_search, gmail_poll,
calendar_poll, send_email, create_task). Results are fed back for final response.
"""

import json
import logging
from datetime import datetime, timezone
from html import escape as html_escape

import os
import re as _re
import xml.etree.ElementTree as ET

import httpx

from feldwebel import get_ctx
from feldwebel.history import add_message, get_history, trim_history

# ── Economic data API keys ──
FRED_API_KEY = os.environ.get("FRED_API_KEY", "")

logger = logging.getLogger("feldwebel.responder")

WEEKDAYS_HU = ["hétfő", "kedd", "szerda", "csütörtök", "péntek", "szombat", "vasárnap"]

# ── Tools definition for DeepSeek ──────────────────────────────────

TOOLS = [
    # ── Email ──
    {"type": "function", "function": {
        "name": "gmail_search", "description": "Keress emaileket a Gmail-ben. Személynév, tárgy, dátum alapján.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string", "description": "Gmail query: from:Darvas, subject:EdTech, newer_than:1d"}}, "required": ["query"]}}},
    {"type": "function", "function": {
        "name": "gmail_poll", "description": "Gmail inbox frissítése — új emailek lekérése.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "send_email", "description": "Email küldés. MINDIG kérj megerősítést a Kommandanttól küldés előtt!",
        "parameters": {"type": "object", "properties": {
            "to": {"type": "string", "description": "Címzett email cím"},
            "subject": {"type": "string", "description": "Email tárgya"},
            "body": {"type": "string", "description": "Email szövege. Aláírás: Üdvözlettel, Kende Tamás"},
        }, "required": ["to", "subject", "body"]}}},
    {"type": "function", "function": {
        "name": "read_gmail_attachment", "description": "Email csatolmány tartalmának kiolvasása (docx, pdf, txt). Először gmail_search-csel keresd meg az email message_id-ját.",
        "parameters": {"type": "object", "properties": {
            "message_id": {"type": "string", "description": "Gmail message ID (gmail_search eredményéből)"},
            "attachment_index": {"type": "integer", "description": "Hanyadik csatolmány (0 = első)", "default": 0},
        }, "required": ["message_id"]}}},
    # ── Naptár ──
    {"type": "function", "function": {
        "name": "calendar_poll", "description": "Naptár frissítése — mai események lekérése.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "calendar_create", "description": "Új naptáresemény létrehozása. MINDIG kérj megerősítést!",
        "parameters": {"type": "object", "properties": {
            "summary": {"type": "string", "description": "Esemény neve"},
            "start_time": {"type": "string", "description": "Kezdés ISO formátumban: 2026-04-10T15:00:00"},
            "end_time": {"type": "string", "description": "Befejezés ISO formátumban: 2026-04-10T16:00:00"},
            "description": {"type": "string", "description": "Részletek (opcionális)", "default": ""},
        }, "required": ["summary", "start_time", "end_time"]}}},
    # ── Feladatok ──
    {"type": "function", "function": {
        "name": "create_task", "description": "Új Bridge feladat létrehozása.",
        "parameters": {"type": "object", "properties": {"title": {"type": "string", "description": "Feladat leírása"}}, "required": ["title"]}}},
    {"type": "function", "function": {
        "name": "list_tasks", "description": "Nyitott feladatok listázása.",
        "parameters": {"type": "object", "properties": {}}}},
    # ── AI agentek ──
    {"type": "function", "function": {
        "name": "ai_query", "description": "Feladat EGY agentnek: kimi (kutató), deepseek (elemző), glm5 (kódoló/MCP specialist). Web search képes!",
        "parameters": {"type": "object", "properties": {
            "model": {"type": "string", "description": "kimi / deepseek / glm5", "enum": ["kimi", "deepseek", "glm5"]},
            "prompt": {"type": "string", "description": "Feladat / kérdés"},
        }, "required": ["model", "prompt"]}}},
    {"type": "function", "function": {
        "name": "ai_task", "description": "Feladat MIND A 3 agentnek párhuzamosan. Komplex kutatáshoz, elemzéshez.",
        "parameters": {"type": "object", "properties": {
            "title": {"type": "string", "description": "Rövid cím"},
            "description": {"type": "string", "description": "Részletes leírás"},
        }, "required": ["title", "description"]}}},
    {"type": "function", "function": {
        "name": "read_task_results", "description": "AI feladat eredményeinek lekérése. 'Mi lett a feladattal?', 'kész van?'",
        "parameters": {"type": "object", "properties": {"task_id": {"type": "integer", "description": "Task ID szám"}}, "required": ["task_id"]}}},
    # ── Web ──
    {"type": "function", "function": {
        "name": "web_fetch", "description": "Weboldal tartalmának letöltése URL alapján. Ha linket kaptál és többet akarsz olvasni.",
        "parameters": {"type": "object", "properties": {"url": {"type": "string", "description": "A letöltendő URL"}}, "required": ["url"]}}},
    # ── Bridge memória és kommunikáció ──
    {"type": "function", "function": {
        "name": "read_memory", "description": "Bridge shared memory olvasása. Korábbi döntések, tudás, kontextus.",
        "parameters": {"type": "object", "properties": {
            "key": {"type": "string", "description": "Memória kulcs (opcionális)", "default": ""},
            "category": {"type": "string", "description": "Kategória szűrés (opcionális)", "default": ""},
        }}}},
    {"type": "function", "function": {
        "name": "search_memory", "description": "Keresés a Bridge shared memory-ban kulcsszó alapján.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string", "description": "Keresési kifejezés"}}, "required": ["query"]}}},
    {"type": "function", "function": {
        "name": "send_message", "description": "Bridge üzenet küldés másik instance-nak (web-claus, cli-claus, stb.).",
        "parameters": {"type": "object", "properties": {
            "recipient": {"type": "string", "description": "Címzett: web-claus, cli-claus, kommandant"},
            "subject": {"type": "string", "description": "Tárgy"},
            "message": {"type": "string", "description": "Üzenet szövege"},
        }, "required": ["recipient", "subject", "message"]}}},
    {"type": "function", "function": {
        "name": "search_discussions", "description": "Keresés a Bridge diskussziókban. Korábbi viták, döntések.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string", "description": "Keresési kifejezés"}}, "required": ["query"]}}},
    # ── Recipe-k (Operation Zahnrad) ──
    {"type": "function", "function": {
        "name": "list_recipes", "description": "Elérhető recipe-k (workflow template-ek) listázása.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "execute_recipe", "description": "Recipe futtatása — workflow végrehajtás AI agenten keresztül.",
        "parameters": {"type": "object", "properties": {
            "name": {"type": "string", "description": "Recipe neve (pl. daily_briefing, weekly_macro_report)"},
            "model": {"type": "string", "description": "Agent: kimi / deepseek / glm5", "default": "deepseek"},
        }, "required": ["name"]}}},
    {"type": "function", "function": {
        "name": "create_recipe", "description": "Új recipe létrehozása — újrahasználható workflow template.",
        "parameters": {"type": "object", "properties": {
            "name": {"type": "string", "description": "Egyedi recipe név"},
            "description": {"type": "string", "description": "Mit csinál a recipe"},
            "prompt_template": {"type": "string", "description": "A workflow prompt szövege"},
        }, "required": ["name", "description", "prompt_template"]}}},
    # ── Cron ütemezés ──
    {"type": "function", "function": {
        "name": "schedule_recipe", "description": "Recipe ütemezése — automatikus futtatás cron időzítéssel. Eredmény Telegramra jön.",
        "parameters": {"type": "object", "properties": {
            "name": {"type": "string", "description": "Recipe neve"},
            "cron_schedule": {"type": "string", "description": "Cron: '0 7 * * *' (7:00 minden nap), '0 8 * * 1' (hétfő 8:00), '30 17 * * 1-5' (hétköznap 17:30). 'off' = kikapcsol"},
            "model": {"type": "string", "description": "Agent: glm5 / deepseek / kimi / all", "default": "glm5"},
        }, "required": ["name", "cron_schedule"]}}},
    # ── Fájl export ──
    {"type": "function", "function": {
        "name": "export_task", "description": "AI feladat exportálása fájlba. Telegramra küld, és opcionálisan emailben is csatolmányként.",
        "parameters": {"type": "object", "properties": {
            "task_id": {"type": "integer", "description": "AI task ID száma"},
            "format": {"type": "string", "description": "Formátum: docx / xlsx / pptx", "enum": ["docx", "xlsx", "pptx"]},
            "email_to": {"type": "string", "description": "Ha emailben is kell küldeni, ide írd a címet", "default": ""},
        }, "required": ["task_id", "format"]}}},
    # ── Fájl továbbítás ──
    {"type": "function", "function": {
        "name": "send_file_email", "description": "Feltöltött fájl (fotó, dokumentum) továbbküldése emailben. Upload #ID szükséges.",
        "parameters": {"type": "object", "properties": {
            "file_id": {"type": "integer", "description": "Upload fájl ID (pl. #2)"},
            "to": {"type": "string", "description": "Email cím"},
            "subject": {"type": "string", "description": "Email tárgy", "default": "Claus — csatolt fájl"},
        }, "required": ["file_id", "to"]}}},
    # ── Kép elemzés ──
    {"type": "function", "function": {
        "name": "analyze_upload", "description": "Feltöltött kép elemzése Kimi K2.6 vision modellel. Upload #ID szükséges.",
        "parameters": {"type": "object", "properties": {
            "file_id": {"type": "integer", "description": "Upload fájl ID (pl. #2)"},
            "prompt": {"type": "string", "description": "Kérdés a képről", "default": "Mit latsz a kepen? Ird le reszletesen, magyarul."},
        }, "required": ["file_id"]}}},
    # ── Közgazdasági adatok (Makronóm) ──
    {"type": "function", "function": {
        "name": "mnb_rates", "description": "ECB napi árfolyamok (EUR/HUF, USD/HUF, CHF/HUF stb.) + HUF keresztárfolyamok. Friss, mai adat.",
        "parameters": {"type": "object", "properties": {
            "currencies": {"type": "string", "description": "Szűrés: 'HUF,USD,CHF' (üres = fő devizák)", "default": ""},
        }}}},
    {"type": "function", "function": {
        "name": "market_quote", "description": "Piaci adat Yahoo Finance-ből: részvény, deviza, áru, index, kötvényhozam, kripto. FRISS árak.",
        "parameters": {"type": "object", "properties": {
            "symbol": {"type": "string", "description": "Ticker: EURHUF=X, OTP.BD, ^BUX.BD, ^GSPC, GC=F (arany), BZ=F (Brent), BTC-USD, ^TNX (US 10Y)"},
        }, "required": ["symbol"]}}},
    {"type": "function", "function": {
        "name": "fred_data", "description": "FRED makro idősorok: US kamatok, infláció, GDP, munkanélküliség, hozamok, M2, VIX, stb.",
        "parameters": {"type": "object", "properties": {
            "series_id": {"type": "string", "description": "FRED ID: UNRATE, DGS10, DGS2, T10Y2Y, CPIAUCSL, GDP, DFF, MORTGAGE30US, UMCSENT, M2SL, VIXCLS"},
            "limit": {"type": "integer", "description": "Utolsó N adatpont", "default": 12},
        }, "required": ["series_id"]}}},
    {"type": "function", "function": {
        "name": "policy_rates", "description": "Jegybanki alapkamatok — MNB, ECB, Fed, CNB, NBP, BNR, BoE, stb. BIS/DBnomics forrás.",
        "parameters": {"type": "object", "properties": {
            "countries": {"type": "string", "description": "BIS kódok: HU (MNB), XM (ECB), US (Fed), CZ, PL, RO, GB, JP, CH, TR", "default": "HU,XM,US,CZ,PL"},
        }}}},
    {"type": "function", "function": {
        "name": "econ_calendar", "description": "Közelgő gazdasági adatközlések és jegybanki események (US + EU).",
        "parameters": {"type": "object", "properties": {
            "days_ahead": {"type": "integer", "description": "Hány napra előre (max 30)", "default": 7},
        }}}},
]

# Tools that need confirmation before execution
CONFIRM_TOOLS = {"send_email"}  # Future: add more dangerous actions


# ── System prompt ──────────────────────────────────────────────────

SYSTEM_PROMPT_TEMPLATE = """Te a Feldwebel vagy — a Kommandant személyes Telegram asszisztense a Claus-Bridge rendszeren belül.

Mai dátum: {date_str} ({weekday_hu})
Időzóna: CET (Budapest)

Szereped:
- Gyors, lényegre törő válaszok magyarul
- Tool-okat tudsz hívni: Gmail keresés, naptár, feladat kezelés, AI agentek
- Ha a Kommandant személyt vagy emailt keres, MINDIG használd a gmail_search tool-t
- Ha frissítést kér, használd a gmail_poll vagy calendar_poll tool-t

FONTOS — AI agentek (rizsrakéták):
- **Kimi** (model="kimi"): 256k kontextus, kutató/elemző
- **DeepSeek** (model="deepseek"): gyors gondolkodás, kritikus
- **GLM-5.1** (model="glm5"): 200k kontextus, 128k output, kódoló + MCP tool-use specialist

SZABÁLYOK az agentekhez:
- Ha a Kommandant azt mondja "adj ki feladatot", "kérdezd meg", "nézzen utána", "elemezze", vagy BÁRMILYEN agentet említ név szerint → AZONNAL hívd a megfelelő tool-t! NE keress Gmail-ben, NE válaszolj szövegben!
- Az agentek önálló AI modellek, NEM email kontaktok — SOHA ne keress rájuk a Gmail-ben!
- **ai_query**: egyszerű kérdés egy agentnek (NEM tud webet keresni, csak a tudásából válaszol)
- **ai_task**: kutatás, elemzés, aktuális adatok keresése — WEB SEARCH KÉPES! Az agentek DuckDuckGo-val keresnek.
- Ha AKTUÁLIS/FRISS adatokat kell keresni (közvélemény-kutatás, árak, hírek) → MINDIG ai_task-ot használj, NE ai_query-t!
- KÖZGAZDASÁGI ADATOKHOZ van 5 dedikált tool — EZEKET HASZNÁLD, NE ai_task-ot!
  * mnb_rates — MNB hivatalos árfolyamok (EUR/HUF, USD/HUF stb.)
  * market_quote — részvény, deviza, árupiaci árak Yahoo Finance-ből (EURHUF=X, OTP.BD, ^BUX.BD, GC=F, BZ=F, BTC-USD)
  * fred_data — US makro: kamatok (DGS10, DFF), infláció (CPIAUCSL), GDP, munkanélküliség (UNRATE), VIX stb.
  * policy_rates — jegybanki alapkamatok (MNB, ECB, Fed, CNB, NBP stb.)
  * econ_calendar — közelgő adatközlések és jegybanki események
- Ha a Kommandant árfolyamot, kamatot, piaci adatot kér → AZONNAL hívd a megfelelő közgazdasági tool-t!
- "Adj ki feladatot Kiminek kutatásra" → ai_task(title="...", description="Kimi feladata: ...")
- "Kérdezd meg GLM-5.1-et mi a véleménye" → ai_query(model="glm5", prompt="...")
- "Mind a hárman elemezzétek" → ai_task(title="...", description="...")

RECIPE-K (workflow template-ek) — FONTOS:
- A Bridge-en vannak előre definiált workflow-k ("recipe-k"). Egy recipe = egy előre megírt prompt, amit EGY agent hajt végre.
- MINDIG az execute_recipe tool-t használd recipe futtatásra! SOHA NE használj ai_task-ot recipe-hez!
- Egy agent: execute_recipe(name="...", model="glm5") — gyors, ~30mp
- Mind a 3 agent: execute_recipe(name="...", model="all") — alapos de lassabb (~3-5 perc)
- Ha a Kommandant konkrét agentet kér (pl. "GLM-5.1-gyel"): model="glm5"
- Ha "mind a hárman" / "összes agent" / "alaposan": model="all"
- Alapértelmezés: model="deepseek" (gyors, egy agent)
- "Futtasd le a napi briefet" / "reggeli brief" → execute_recipe(name="daily_briefing")
- "Napi hírek" / "hírösszefoglaló" / "news brief" → execute_recipe(name="daily_news_brief")
- "Heti makro riport" → execute_recipe(name="weekly_macro_report")
- "Milyen recipe-k vannak?" → list_recipes()
- "Csináljatok egy új recipe-t..." → create_recipe(name="...", description="...", prompt_template="...")
- Ha a Kommandant egy ismétlődő feladatot kér és még nincs rá recipe → JAVASOLJ recipe létrehozást!

ÜTEMEZÉS (cron):
- "Futtasd a napi briefet minden reggel 7-kor" → schedule_recipe(name="daily_news_brief", cron_schedule="0 7 * * *")
- "Heti makro riport hétfőn 8-kor" → schedule_recipe(name="weekly_macro_report", cron_schedule="0 8 * * 1", model="all")
- "Kapcsold ki az ütemezést" → schedule_recipe(name="...", cron_schedule="off")
- Az ütemezett recipe-k AUTOMATIKUSAN futnak és Telegramra küldik az eredményt

FOTÓK ÉS FÁJLOK:
- A Kommandant Telegramon küldhet fotókat — azok upload #ID-val tárolódnak
- "Elemezd a #2-es képet" → analyze_upload(file_id=2)
- "Küldd el emailben a #2-es fotót a cim@email.com-ra" → send_file_email(file_id=2, to="cim@email.com")
- Email küldéshez MINDIG kérj megerősítést!

FÁJL EXPORT (docx/xlsx/pptx):
- "Küld el docx-ben" / "excelt kérek" / "pptx-et a task #45-ről" → export_task(task_id=45, format="pptx")
- Ha a Kommandant nem ad task_id-t, kérdezd meg melyik taskra gondol, vagy használd az utolsó completed taskot
- A fájl AUTOMATIKUSAN megérkezik Telegramon
- Ha emailben is kell: export_task(task_id=45, format="docx", email_to="cim@example.com")
- Email küldéshez MINDIG kérj megerősítést a Kommandanttól MIELŐTT meghívod a tool-t!

Ha emailre kell válaszolni:
- Írd meg a draft választ magyarul, a Kommandant stílusában (hivatalos de nem merev)
- Kérd a megerősítését mielőtt bármit küldenél
- Aláírás: "Üdvözlettel, Kende Tamás"

Szabályok:
- Magyar nyelv, hacsak nem kérnek mást
- Tömör válaszok (max 3-5 mondat, kivéve ha elemzést kérnek)
- NE HAZUDJ — ha a tool nem ad eredményt, mondd meg
- Email küldés SOHA nem automatikus — mindig kérj megerősítést"""


def _build_system_prompt(ctx) -> str:
    """Build system prompt with current date + Pyramid context."""
    now = datetime.now(timezone.utc)
    base = SYSTEM_PROMPT_TEMPLATE.format(
        date_str=now.strftime("%Y.%m.%d"),
        weekday_hu=WEEKDAYS_HU[now.weekday()],
    )

    parts = [base]

    # Add Pyramid context if available
    try:
        from pyramid.context_builder import build_agent_context
        pyramid_ctx = build_agent_context(
            agent_id="deepseek",
            inbox_summary="",  # Tools handle email access now
        )
        parts.append(pyramid_ctx)
    except (ImportError, Exception) as e:
        logger.debug("Pyramid context unavailable: %s", e)

    # Add open tasks (fast, from DB)
    task_ctx = _fetch_open_tasks(ctx)
    if task_ctx:
        parts.append(task_ctx)

    return "\n\n---\n\n".join(parts)


# ── Main respond function (agent loop) ────────────────────────────


async def respond(text: str, chat_id: str, agent_id: str = "deepseek") -> str:
    """
    Generate a response with native tool use.
    Agent loop: DeepSeek calls tools → results fed back → final response.
    First round: tool_choice=required (forces tool use, no hallucination).
    """
    ctx = get_ctx()

    # Store user message in history
    add_message(chat_id, "user", text)

    # Build messages
    system_prompt = _build_system_prompt(ctx)
    history = get_history(chat_id, limit=8)
    messages = [{"role": "system", "content": system_prompt}]
    for msg in history:
        messages.append({"role": msg["role"], "content": msg["content"]})

    model_id = ctx.siliconflow_models.get(agent_id, "deepseek-ai/DeepSeek-V4-Pro")

    # Agent loop — max 5 rounds of tool calls (last round forces text)
    # Round 0: force_tool=True → DeepSeek MUST call a tool (no hallucination)
    final_content = ""
    for round_num in range(5):
        data = await _call_deepseek(
            ctx, model_id, messages,
            use_tools=(round_num < 4),
            force_tool=(round_num == 0),
        )
        if not data:
            break

        choice = data.get("choices", [{}])[0]
        msg = choice.get("message", {})
        finish_reason = choice.get("finish_reason", "")

        content = msg.get("content", "") or ""
        # Safety net: if a thinking-mode SF response delivered the actual
        # answer in reasoning_content with empty content, fall back to it
        # rather than treating the whole round as empty.
        if not content.strip():
            reasoning = msg.get("reasoning_content", "") or ""
            if reasoning.strip():
                content = reasoning
        tool_calls = msg.get("tool_calls")

        # Case A: Proper JSON tool_calls
        if tool_calls and finish_reason == "tool_calls":
            messages.append(msg)
            for tc in tool_calls:
                fn_name = tc.get("function", {}).get("name", "")
                fn_args_raw = tc.get("function", {}).get("arguments", "{}")
                tc_id = tc.get("id", "")
                try:
                    fn_args = json.loads(fn_args_raw) if isinstance(fn_args_raw, str) else fn_args_raw
                except json.JSONDecodeError:
                    fn_args = {}
                logger.info("Tool call round %d: %s(%s)", round_num, fn_name, fn_args)
                result = await _execute_tool(fn_name, fn_args, ctx)
                messages.append({"role": "tool", "tool_call_id": tc_id, "content": result})
            continue

        # Case B: Text-based tool calls (DeepSeek DSML or Kimi markers)
        import re
        text_markers = ["<｜DSML｜", "<|tool_call", "function_calls>", "tool_calls_section"]
        if not tool_calls and any(m in content for m in text_markers):
            # Parse tool name and parameters from text
            # DSML format: <｜DSML｜invoke name="tool_name"> <｜DSML｜parameter name="x" string="true">value</｜DSML｜parameter>
            # Kimi format: <|tool_call_begin|>functions.web_search:2<|tool_call_argument_begin|>{"query": "..."}
            parsed_tool = _parse_text_tool_call(content)
            if parsed_tool:
                fn_name, fn_args = parsed_tool
                logger.info("Text tool call round %d: %s(%s)", round_num, fn_name, fn_args)
                result = await _execute_tool(fn_name, fn_args, ctx)
                messages.append({"role": "assistant", "content": f"(Tool végrehajtva: {fn_name})"})
                messages.append({"role": "user", "content": f"Tool eredmény ({fn_name}):\n{result}\n\nFoglald össze az eredményt a Kommandantnak."})
                continue

        # Text response — we're done
        final_content = content
        break

    if not final_content:
        # Debug: show what happened in each round
        debug_parts = [f"Agent loop ended with no content (model={model_id})"]
        # Show last few messages for debugging
        for m in messages[-4:]:
            role = m.get("role", "?")
            c = (m.get("content") or "")[:150]
            tc = m.get("tool_calls")
            if tc:
                debug_parts.append(f"  [{role}] tool_calls={json.dumps([t.get('function',{}).get('name','?') for t in tc])}")
            else:
                debug_parts.append(f"  [{role}] {c[:150]}")
        debug_msg = "\n".join(debug_parts)
        logger.error("Feldwebel empty response: %s", debug_msg)
        await ctx.telegram_push(f"<b>FELDWEBEL DEBUG</b>\n<pre>{html_escape(debug_msg[:1500])}</pre>")
        return ""

    # Store in history
    add_message(chat_id, "assistant", final_content, agent_id)

    # Push to Telegram
    safe_content = html_escape(final_content[:3800])
    await ctx.telegram_push(f"<b>FELDWEBEL</b>\n\n{safe_content}")

    # Store in Bridge DB
    _store_bridge_reply(ctx, text, final_content, agent_id)

    # Trim old history
    trim_history(chat_id, max_entries=30)

    logger.info("Feldwebel responded (%d chars, tool-use)", len(final_content))
    return final_content


# ── DeepSeek API call ──────────────────────────────────────────────


async def _call_deepseek(ctx, model_id: str, messages: list, use_tools: bool = True, force_tool: bool = False) -> dict:
    """Call SiliconFlow AI API, optionally with tools."""
    try:
        payload = {
            "model": model_id,
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 2000,
        }
        # K2.6 + V4-Pro default thinking ON on SF — clamp HARD in the agent loop.
        # reasoning_effort=medium proved insufficient for V4-Pro: with
        # tool_choice="required" on round 0 the model reasoned but never committed
        # to content OR a tool call (5 empty rounds). The agent loop is purely
        # tool-driven, so thinking buys nothing here — disable it for both.
        if "Kimi" in model_id or "DeepSeek" in model_id:
            payload["thinking"] = {"type": "disabled"}
        if use_tools:
            payload["tools"] = TOOLS
            if force_tool:
                payload["tool_choice"] = "required"

        async with httpx.AsyncClient(timeout=ctx.siliconflow_timeout) as client:
            resp = await client.post(
                f"{ctx.siliconflow_base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {ctx.siliconflow_api_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )
            data = json.loads(resp.text)

        if "error" in data:
            logger.error("SiliconFlow error: %s", data["error"])
            await ctx.telegram_push(f"<b>FELDWEBEL</b> — API hiba: {html_escape(str(data['error'])[:500])}")
            return {}
        return data

    except Exception as e:
        logger.error("SiliconFlow API call failed: %s", e)
        await ctx.telegram_push(f"<b>FELDWEBEL</b> — API hiba: {html_escape(str(e)[:500])}")
        return {}


# ── Tool execution ─────────────────────────────────────────────────


async def _execute_tool(name: str, args: dict, ctx) -> str:
    """Execute a tool and return the result as string."""
    # IMPORTANT: import httpx here to avoid UnboundLocalError —
    # analyze_upload handler has a conditional `import httpx` which makes
    # Python treat httpx as local for the entire function scope.
    import httpx

    if name == "gmail_search":
        query = args.get("query", "")
        if not query:
            return json.dumps({"error": "Nincs keresési query"})
        if not ctx.capture_state.get("_capture_inbox_func"):
            return json.dumps({"error": "Gmail nem elérhető"})
        try:
            result = await ctx.capture_state["_capture_inbox_func"](
                limit=10, query=query, caller="feldwebel"
            )
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "gmail_poll":
        if not ctx.capture_gmail_poll:
            return json.dumps({"error": "Gmail poll nem elérhető"})
        try:
            result = await ctx.capture_gmail_poll(caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "calendar_poll":
        if not ctx.capture_calendar_poll:
            return json.dumps({"error": "Calendar poll nem elérhető"})
        try:
            result = await ctx.capture_calendar_poll(caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "create_task":
        title = args.get("title", "")
        if not title:
            return json.dumps({"error": "Nincs feladat megadva"})
        try:
            conn = ctx.get_db()
            conn.execute(
                "INSERT INTO tasks (title, assigned_to, assigned_by, status, priority, created_at) "
                "VALUES (?, 'cli-claus', 'feldwebel', 'pending', 'normal', ?)",
                (title, datetime.now(timezone.utc).isoformat())
            )
            conn.commit()
            conn.close()
            return json.dumps({"success": True, "title": title})
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "list_tasks":
        try:
            conn = ctx.get_db()
            rows = conn.execute(
                "SELECT id, title, status, priority, assigned_to FROM tasks "
                "WHERE status IN ('pending', 'in_progress') "
                "ORDER BY created_at DESC LIMIT 15"
            ).fetchall()
            conn.close()
            tasks = [{"id": r["id"], "title": r["title"], "status": r["status"],
                      "priority": r["priority"], "assigned_to": r["assigned_to"]} for r in rows]
            return json.dumps({"tasks": tasks}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "ai_query":
        model = args.get("model", "deepseek")
        prompt = args.get("prompt", "")
        if not prompt:
            return json.dumps({"error": "Nincs prompt megadva"})

        # Route through ai_task for web search capability
        ai_task_func = ctx.capture_state.get("_ai_task_func")
        if ai_task_func:
            try:
                now = datetime.now(timezone.utc)
                date_info = f"\n\n[Mai dátum: {now.strftime('%Y.%m.%d')} ({WEEKDAYS_HU[now.weekday()]})]"
                agent_tasks = json.dumps({model: {"prompt": prompt + date_info, "max_tokens": 3000}})
                result_json = await ai_task_func(
                    title=f"Feldwebel → {model}: {prompt[:60]}",
                    description=prompt + date_info,
                    assigned_by="feldwebel",
                    agent_tasks=agent_tasks,
                )
                result = json.loads(result_json)
                task_id = result.get("task_id")

                if task_id:
                    # Poll for result (ai_task runs in background thread)
                    import asyncio
                    for _ in range(60):  # max 120s wait
                        await asyncio.sleep(2)
                        conn = ctx.get_db()
                        row = conn.execute(
                            "SELECT content FROM ai_task_results WHERE task_id = ? LIMIT 1",
                            (task_id,)
                        ).fetchone()
                        status = conn.execute(
                            "SELECT status FROM ai_tasks WHERE id = ?", (task_id,)
                        ).fetchone()
                        conn.close()
                        if row:
                            return json.dumps({
                                "model": model, "response": row["content"],
                                "task_id": task_id
                            }, ensure_ascii=False)
                        if status and status["status"] == "failed":
                            return json.dumps({"error": f"AI task #{task_id} failed"})
                    return json.dumps({"error": f"AI task #{task_id} timeout (120s)"})
                return result_json
            except Exception as e:
                logger.error("ai_query via ai_task failed: %s", e)

        # Fallback: direct ai_query (no web search)
        ai_query_func = ctx.capture_state.get("_ai_query_func")
        if ai_query_func:
            try:
                return await ai_query_func(model=model, prompt=prompt, caller="feldwebel")
            except Exception as e:
                return json.dumps({"error": str(e)})

        return json.dumps({"error": "AI query nem elérhető"})

    if name == "ai_task":
        title = args.get("title", "")
        description = args.get("description", "")
        if not title or not description:
            return json.dumps({"error": "Cím és leírás szükséges"})
        if not ctx.capture_state.get("_ai_task_func"):
            return json.dumps({"error": "AI task nem elérhető"})
        try:
            now = datetime.now(timezone.utc)
            date_info = f"\n\n[Mai dátum: {now.strftime('%Y.%m.%d')} ({WEEKDAYS_HU[now.weekday()]}). Az adatoknak FRISSNEK kell lenniük!]"
            result_json = await ctx.capture_state["_ai_task_func"](
                title=title, description=description + date_info, assigned_by="feldwebel"
            )
            # Return immediately — user can ask for results later with read_task_results
            return result_json
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "read_task_results":
        task_id = args.get("task_id", 0)
        if not task_id:
            return json.dumps({"error": "task_id szükséges"})
        try:
            conn = ctx.get_db()
            task = conn.execute(
                "SELECT title, status FROM ai_tasks WHERE id = ?", (task_id,)
            ).fetchone()
            if not task:
                conn.close()
                return json.dumps({"error": f"Task #{task_id} nem található"})
            if task["status"] in ("pending", "running"):
                conn.close()
                return json.dumps({"task_id": task_id, "status": task["status"],
                                   "message": f"Task #{task_id} még fut, az agentek dolgoznak."})
            rows = conn.execute(
                "SELECT agent, content FROM ai_task_results WHERE task_id = ? ORDER BY timestamp",
                (task_id,)
            ).fetchall()
            conn.close()
            if rows:
                parts = []
                for r in rows:
                    parts.append(f"**{r['agent'].upper()}**:\n{r['content'][:2000]}")
                return json.dumps({
                    "task_id": task_id, "title": task["title"], "status": task["status"],
                    "agent_results": "\n\n---\n\n".join(parts)
                }, ensure_ascii=False)
            return json.dumps({"task_id": task_id, "status": task["status"], "message": "Nincs eredmény"})
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "send_email":
        to = args.get("to", "")
        subject = args.get("subject", "")
        body = args.get("body", "")
        if not to or not subject or not body:
            return json.dumps({"error": "to, subject, body mind szükséges"})
        send_func = ctx.capture_state.get("_send_email_func")
        if not send_func:
            return json.dumps({"error": "Email küldés nem elérhető"})
        try:
            result = await send_func(to=to, subject=subject, body=body, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "read_gmail_attachment":
        message_id = args.get("message_id", "")
        attachment_index = args.get("attachment_index", 0)
        if not message_id:
            return json.dumps({"error": "message_id szükséges (gmail_search-ből)"})
        read_att_func = ctx.capture_state.get("_read_gmail_attachment_func")
        if not read_att_func:
            return json.dumps({"error": "Csatolmány olvasás nem elérhető"})
        try:
            result = await read_att_func(message_id=message_id, attachment_index=attachment_index, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "calendar_create":
        summary = args.get("summary", "")
        start = args.get("start_time", "") or args.get("start", "")
        end = args.get("end_time", "") or args.get("end", "")
        description = args.get("description", "")
        if not summary or not start:
            return json.dumps({"error": "summary és start_time szükséges"})
        create_cal_func = ctx.capture_state.get("_create_calendar_func")
        if not create_cal_func:
            return json.dumps({"error": "Naptár létrehozás nem elérhető"})
        try:
            result = await create_cal_func(summary=summary, start=start, end=end, description=description, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "web_fetch":
        url = args.get("url", "")
        if not url or not url.startswith("http"):
            return json.dumps({"error": "Érvényes URL szükséges"})
        import re
        try:
            async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
                resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36"})
            text = resp.text
            text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            return json.dumps({"url": url, "content": text[:4000], "length": len(text)}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "read_memory":
        key = args.get("key", "")
        category = args.get("category", "")
        read_mem_func = ctx.capture_state.get("_read_memory_func")
        if not read_mem_func:
            return json.dumps({"error": "Memory olvasás nem elérhető"})
        try:
            result = await read_mem_func(key=key, category=category, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "search_memory":
        query = args.get("query", "")
        if not query:
            return json.dumps({"error": "Keresési kifejezés szükséges"})
        search_mem_func = ctx.capture_state.get("_search_memory_func")
        if not search_mem_func:
            return json.dumps({"error": "Memory keresés nem elérhető"})
        try:
            result = await search_mem_func(query=query, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "send_message":
        recipient = args.get("recipient", "")
        subject = args.get("subject", "")
        message = args.get("message", "")
        if not recipient or not message:
            return json.dumps({"error": "recipient és message szükséges"})
        send_msg_func = ctx.capture_state.get("_send_message_func")
        if not send_msg_func:
            return json.dumps({"error": "Bridge üzenet küldés nem elérhető"})
        try:
            result = await send_msg_func(sender="feldwebel", recipient=recipient, subject=subject or "Feldwebel üzenet", message=message, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "search_discussions":
        query = args.get("query", "")
        if not query:
            return json.dumps({"error": "Keresési kifejezés szükséges"})
        search_disc_func = ctx.capture_state.get("_search_discussions_func")
        if not search_disc_func:
            return json.dumps({"error": "Diskusszió keresés nem elérhető"})
        try:
            result = await search_disc_func(query=query, caller="feldwebel")
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    # ── Recipe tools (Operation Zahnrad) ──
    if name == "list_recipes":
        try:
            conn = ctx.get_db()
            rows = conn.execute(
                "SELECT name, description, required_tools FROM pyramid_recipes WHERE enabled = 1 ORDER BY name"
            ).fetchall()
            conn.close()
            recipes = [{"name": r["name"], "description": r["description"],
                        "tools": json.loads(r["required_tools"]) if r["required_tools"] else []} for r in rows]
            return json.dumps({"recipes": recipes}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "execute_recipe":
        recipe_name = args.get("name", "")
        model = args.get("model", "deepseek")
        if not recipe_name:
            return json.dumps({"error": "Recipe neve szükséges"})
        try:
            conn = ctx.get_db()
            row = conn.execute(
                "SELECT prompt_template, required_tools, enabled FROM pyramid_recipes WHERE name = ?",
                (recipe_name,)
            ).fetchone()
            conn.close()
            if not row:
                return json.dumps({"error": f"Recipe '{recipe_name}' nem található"})
            if not row["enabled"]:
                return json.dumps({"error": f"Recipe '{recipe_name}' le van tiltva"})
            # Execute via ai_query (dispatch mode for web search support)
            ai_task_func = ctx.capture_state.get("_ai_task_func")
            if not ai_task_func:
                return json.dumps({"error": "AI task nem elérhető"})
            prompt = row["prompt_template"]
            now = datetime.now(timezone.utc)
            date_info = f"\n\n[Mai dátum: {now.strftime('%Y.%m.%d')} ({WEEKDAYS_HU[now.weekday()]})]"
            agent_tasks = json.dumps({model: {"prompt": prompt + date_info, "max_tokens": 8000}})
            result_json = await ai_task_func(
                title=f"Recipe: {recipe_name}",
                description=prompt + date_info,
                assigned_by="feldwebel",
                agent_tasks=agent_tasks,
            )
            result = json.loads(result_json)
            task_id = result.get("task_id")
            if task_id:
                import asyncio
                for _ in range(60):
                    await asyncio.sleep(2)
                    conn = ctx.get_db()
                    row2 = conn.execute("SELECT content FROM ai_task_results WHERE task_id = ? LIMIT 1", (task_id,)).fetchone()
                    status = conn.execute("SELECT status FROM ai_tasks WHERE id = ?", (task_id,)).fetchone()
                    conn.close()
                    if row2:
                        return json.dumps({"recipe": recipe_name, "model": model, "result": row2["content"], "task_id": task_id}, ensure_ascii=False)
                    if status and status["status"] == "failed":
                        return json.dumps({"error": f"Recipe task #{task_id} failed"})
                return json.dumps({"error": f"Recipe task #{task_id} timeout"})
            return result_json
        except Exception as e:
            return json.dumps({"error": str(e)})

    if name == "create_recipe":
        rname = args.get("name", "")
        desc = args.get("description", "")
        prompt_tpl = args.get("prompt_template", "")
        if not rname or not desc or not prompt_tpl:
            return json.dumps({"error": "name, description, prompt_template mind szükséges"})
        try:
            conn = ctx.get_db()
            ts = datetime.now(timezone.utc).isoformat()
            conn.execute(
                "INSERT INTO pyramid_recipes (name, description, required_tools, prompt_template, created_by, created_at, updated_at) "
                "VALUES (?, ?, '[]', ?, 'feldwebel', ?, ?)",
                (rname, desc, prompt_tpl, ts, ts),
            )
            conn.commit()
            conn.close()
            return json.dumps({"status": "created", "name": rname, "message": f"Recipe '{rname}' létrehozva!"}, ensure_ascii=False)
        except Exception as e:
            if "UNIQUE constraint" in str(e):
                return json.dumps({"error": f"Recipe '{rname}' már létezik."})
            return json.dumps({"error": str(e)})

    # ── Cron scheduling ──
    if name == "schedule_recipe":
        recipe_name = args.get("name", "")
        cron_schedule = args.get("cron_schedule", "")
        model = args.get("model", "glm5")
        if not recipe_name or not cron_schedule:
            return json.dumps({"error": "name és cron_schedule szükséges"})
        try:
            conn = ctx.get_db()
            row = conn.execute("SELECT id, name FROM pyramid_recipes WHERE name = ?", (recipe_name,)).fetchone()
            if not row:
                conn.close()
                return json.dumps({"error": f"Recipe '{recipe_name}' nem található"})

            if cron_schedule.lower() == "off":
                conn.execute("UPDATE pyramid_recipes SET cron_enabled = 0 WHERE name = ?", (recipe_name,))
                conn.commit()
                conn.close()
                return json.dumps({"status": "disabled", "name": recipe_name, "message": f"'{recipe_name}' ütemezés kikapcsolva."})

            parts = cron_schedule.strip().split()
            if len(parts) != 5:
                conn.close()
                return json.dumps({"error": "Cron formátum: 'perc óra nap hónap hétnap', pl. '0 7 * * *'"})

            conn.execute(
                "UPDATE pyramid_recipes SET cron_schedule = ?, cron_model = ?, cron_enabled = 1, cron_delivery = 'both' WHERE name = ?",
                (cron_schedule.strip(), model, recipe_name),
            )
            conn.commit()
            conn.close()
            return json.dumps({
                "status": "scheduled", "name": recipe_name,
                "cron_schedule": cron_schedule, "model": model,
                "message": f"'{recipe_name}' ütemezve: {cron_schedule} ({model}). Eredmény Telegramra jön."
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    # ── Send uploaded file via email ──
    if name == "send_file_email":
        fid = args.get("file_id", 0)
        to = args.get("to", "")
        subject = args.get("subject", "Claus — csatolt fajl")
        if not fid or not to:
            return json.dumps({"error": "file_id és to szükséges"})
        send_email_func = ctx.capture_state.get("_send_email_func")
        if not send_email_func:
            return json.dumps({"error": "Email küldés nem elérhető"})
        try:
            result = await send_email_func(
                to=to, subject=subject,
                body=f"Csatolva: fájl #{fid}\n\nKüldte: Claus Bridge",
                file_id=fid, caller="feldwebel",
            )
            return result
        except Exception as e:
            return json.dumps({"error": str(e)})

    # ── Analyze uploaded image ──
    if name == "analyze_upload":
        fid = args.get("file_id", 0)
        prompt = args.get("prompt", "Mit latsz a kepen? Ird le reszletesen, magyarul.")
        if not fid:
            return json.dumps({"error": "file_id szükséges"})
        try:
            conn = ctx.get_db()
            row = conn.execute("SELECT filename, mime_type, content_base64 FROM uploads WHERE id = ?", (fid,)).fetchone()
            conn.close()
            if not row:
                return json.dumps({"error": f"Upload #{fid} nem található"})
            if not row["content_base64"]:
                return json.dumps({"error": f"Upload #{fid} nem tartalmaz képet"})
            if not row["mime_type"].startswith("image/"):
                return json.dumps({"error": f"Upload #{fid} nem kép ({row['mime_type']})"})
            # Use the centralized vision helper via capture_state
            import httpx
            sf_key = ctx.siliconflow_api_key
            sf_base = ctx.siliconflow_base_url
            sf_models = ctx.siliconflow_models
            model_id = sf_models.get("kimi", "moonshotai/Kimi-K2.6")
            data_url = f"data:{row['mime_type']};base64,{row['content_base64']}"
            vision_payload = {
                "model": model_id,
                "messages": [{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": data_url}},
                    {"type": "text", "text": prompt},
                ]}],
                "temperature": 0.5,
                "max_tokens": 2000,
            }
            if "Kimi" in model_id:
                vision_payload["thinking"] = {"type": "disabled"}
            async with httpx.AsyncClient(timeout=60) as client:
                resp = await client.post(
                    f"{sf_base}/chat/completions",
                    headers={"Authorization": f"Bearer {sf_key}", "Content-Type": "application/json"},
                    json=vision_payload,
                )
            data = json.loads(resp.text)
            analysis = data.get("choices", [{}])[0].get("message", {}).get("content", "(ures valasz)")
            return json.dumps({"status": "analyzed", "file_id": fid, "analysis": analysis}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    # ── File export + Telegram send ──
    if name == "export_task":
        task_id = args.get("task_id", 0)
        fmt = args.get("format", "docx")
        if not task_id:
            return json.dumps({"error": "task_id szükséges"})
        if fmt not in ("docx", "xlsx", "pptx"):
            return json.dumps({"error": "format: docx / xlsx / pptx"})
        try:
            conn = ctx.get_db()
            task = conn.execute("SELECT * FROM ai_tasks WHERE id = ?", (task_id,)).fetchone()
            if not task:
                conn.close()
                return json.dumps({"error": f"Task #{task_id} nem található"})
            results = conn.execute(
                "SELECT agent, role, content, timestamp FROM ai_task_results WHERE task_id = ? ORDER BY id",
                (task_id,)
            ).fetchall()
            conn.close()
            if not results:
                return json.dumps({"error": f"Task #{task_id} nincs eredménye"})

            # Generate file
            if fmt == "xlsx":
                gen_func = ctx.capture_state.get("_generate_xlsx")
                if not gen_func:
                    return json.dumps({"error": "xlsx generálás nem elérhető"})
                buf = gen_func(task, results)
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            elif fmt == "pptx":
                gen_func = ctx.capture_state.get("_generate_pptx")
                if not gen_func:
                    return json.dumps({"error": "pptx generálás nem elérhető"})
                buf = gen_func(task, results)
                mime = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
            else:  # docx
                from docx import Document as DocxDocument
                from io import BytesIO
                doc = DocxDocument()
                doc.add_heading(task["title"], level=1)
                doc.add_paragraph(f'Kiadta: {task["assigned_by"]} | Dátum: {task["created_at"][:10]}')
                agent_names = {"kimi": "Kimi-K2.6", "deepseek": "DeepSeek V3.2", "glm5": "GLM-5.1", "szintézis": "Szintézis"}
                for r in results:
                    doc.add_heading(agent_names.get(r["agent"], r["agent"]), level=2)
                    for line in (r["content"] or "").split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            filename = f"claus_task_{task_id}.{fmt}"
            file_bytes = buf.getvalue()
            actions = []

            # Send to Telegram
            send_file = ctx.capture_state.get("_telegram_send_file")
            if send_file:
                caption = f"📎 <b>{task['title']}</b>\nTask #{task_id} | {fmt.upper()}"
                sent = await send_file(file_bytes, filename, mime, caption)
                if sent:
                    actions.append("Telegram")

            # Send as email attachment if requested
            email_to = args.get("email_to", "")
            if email_to and "@" in email_to:
                import base64 as b64mod
                send_email_func = ctx.capture_state.get("_send_email_func")
                if send_email_func:
                    b64_content = b64mod.b64encode(file_bytes).decode("ascii")
                    await send_email_func(
                        to=email_to,
                        subject=f"Claus AI Task #{task_id}: {task['title']}",
                        body=f"Csatolva: {filename}\n\nKészítette: Claus Multi-Agent Rendszer",
                        attachment_base64=b64_content,
                        attachment_filename=filename,
                        attachment_mime=mime,
                        caller="feldwebel",
                    )
                    actions.append(f"email ({email_to})")

            if actions:
                return json.dumps({"status": "sent", "filename": filename, "size": len(file_bytes),
                                   "message": f"{fmt.upper()} elküldve: {', '.join(actions)}"}, ensure_ascii=False)
            return json.dumps({"error": "Sem Telegram, sem email küldés nem sikerült"})
        except Exception as e:
            return json.dumps({"error": str(e)})

    # ── Közgazdasági adat-toolok (Makronóm) ──────────────────────────
    if name in ("mnb_rates", "market_quote", "fred_data", "policy_rates", "econ_calendar"):
        logger.info("ECON TOOL CALLED: %s args=%s", name, args)

    if name == "mnb_rates":
        currencies_filter = [c.strip().upper() for c in args.get("currencies", "").split(",") if c.strip()]
        try:
            # ECB daily reference rates (EUR-based, includes HUF + 30 currencies)
            async with httpx.AsyncClient(timeout=15) as client:
                resp = await client.get(
                    "https://www.ecb.europa.eu/stats/eurofxref/eurofxref-daily.xml",
                    headers={"User-Agent": "ClausBridge/1.0"},
                )
            root = ET.fromstring(resp.text)
            ns = {"ecb": "http://www.ecb.int/vocabulary/2002-08-01/eurofxref"}
            date_str = ""
            for tc in root.findall(".//ecb:Cube[@time]", ns):
                date_str = tc.attrib["time"]
            rates = []
            eur_huf = None
            for cube in root.findall(".//ecb:Cube[@currency]", ns):
                curr = cube.attrib["currency"]
                rate_val = float(cube.attrib["rate"])
                if curr == "HUF":
                    eur_huf = rate_val
                if currencies_filter and curr not in currencies_filter:
                    continue
                rates.append({"currency": curr, "rate": rate_val})
            # If no filter, just return key currencies
            if not currencies_filter:
                key_currencies = {"HUF", "USD", "CHF", "GBP", "CZK", "PLN", "RON", "JPY", "SEK", "NOK", "DKK", "CNY", "TRY"}
                rates = [r for r in rates if r["currency"] in key_currencies]
            # Add cross rates (USD/HUF, CHF/HUF etc.) if HUF is in data
            cross_rates = []
            if eur_huf:
                for r in rates:
                    if r["currency"] != "HUF":
                        cross = round(eur_huf / r["rate"], 2)
                        cross_rates.append({"pair": f"{r['currency']}/HUF", "rate": cross})
            return json.dumps({
                "source": "ECB daily reference rates",
                "date": date_str,
                "base": "EUR",
                "rates": rates,
                "huf_cross_rates": cross_rates,
            }, ensure_ascii=False)
        except Exception as e:
            logger.error("mnb_rates tool error: %s: %s", type(e).__name__, e)
            return json.dumps({"error": f"ECB API hiba: {type(e).__name__}: {e}"})

    if name == "market_quote":
        symbol = args.get("symbol", "").strip()
        if not symbol:
            return json.dumps({"error": "symbol szükséges (pl. EURHUF=X)"})
        try:
            url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?interval=1d&range=5d"
            async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
                resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            data = json.loads(resp.text)
            result_data = data.get("chart", {}).get("result", [])
            if not result_data:
                err = data.get("chart", {}).get("error", {})
                return json.dumps({"error": f"Yahoo Finance: {err.get('description', 'no data')}", "symbol": symbol})
            meta = result_data[0].get("meta", {})
            indicators = result_data[0].get("indicators", {}).get("quote", [{}])[0]
            closes = indicators.get("close", [])
            timestamps = result_data[0].get("timestamp", [])
            price = meta.get("regularMarketPrice", 0)
            prev_close = meta.get("previousClose") or meta.get("chartPreviousClose", 0)
            change = round(price - prev_close, 4) if prev_close else 0
            change_pct = round((change / prev_close) * 100, 2) if prev_close else 0
            # Last 5 closes
            recent = []
            for i in range(max(0, len(closes) - 5), len(closes)):
                if i < len(timestamps) and closes[i] is not None:
                    from datetime import datetime as _dt
                    d = _dt.utcfromtimestamp(timestamps[i]).strftime("%m-%d")
                    recent.append({"date": d, "close": round(closes[i], 4)})
            return json.dumps({
                "symbol": symbol, "name": meta.get("shortName", symbol),
                "price": price, "change": change, "change_pct": change_pct,
                "currency": meta.get("currency", ""),
                "recent_closes": recent,
            }, ensure_ascii=False)
        except Exception as e:
            logger.error("market_quote tool error: %s: %s", type(e).__name__, e)
            return json.dumps({"error": f"Yahoo Finance hiba: {type(e).__name__}: {e}", "symbol": symbol})

    if name == "fred_data":
        series_id = args.get("series_id", "").strip().upper()
        limit = min(args.get("limit", 12), 100)
        if not series_id:
            return json.dumps({"error": "series_id szükséges (pl. UNRATE, DGS10)"})
        if not FRED_API_KEY:
            return json.dumps({"error": "FRED_API_KEY nincs beállítva a Railway env-ben"})
        try:
            url = (
                f"https://api.stlouisfed.org/fred/series/observations"
                f"?series_id={series_id}&api_key={FRED_API_KEY}&file_type=json"
                f"&sort_order=desc&limit={limit}"
            )
            async with httpx.AsyncClient(timeout=15) as client:
                resp = await client.get(url)
            data = json.loads(resp.text)
            if "error_code" in data:
                return json.dumps({"error": data.get("error_message", "FRED error"), "series_id": series_id})
            obs = data.get("observations", [])
            points = [{"date": o["date"], "value": o["value"]} for o in obs if o.get("value") != "."]
            # Also get series info
            info_url = (
                f"https://api.stlouisfed.org/fred/series"
                f"?series_id={series_id}&api_key={FRED_API_KEY}&file_type=json"
            )
            async with httpx.AsyncClient(timeout=10) as client:
                info_resp = await client.get(info_url)
            info_data = json.loads(info_resp.text)
            serieses = info_data.get("seriess", [{}])
            title = serieses[0].get("title", series_id) if serieses else series_id
            units = serieses[0].get("units", "") if serieses else ""
            freq = serieses[0].get("frequency", "") if serieses else ""
            return json.dumps({
                "series_id": series_id, "title": title, "units": units,
                "frequency": freq, "observations": points,
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": f"FRED API hiba: {e}", "series_id": series_id})

    if name == "policy_rates":
        countries = [c.strip().upper() for c in args.get("countries", "HU,XM,US,CZ,PL").split(",") if c.strip()]
        try:
            results = {}
            async with httpx.AsyncClient(timeout=20) as client:
                for cc in countries[:8]:  # max 8 countries
                    url = f"https://api.db.nomics.world/v22/series/BIS/WS_CBPOL/M.{cc}?observations=1&limit=1"
                    resp = await client.get(url, headers={"User-Agent": "ClausBridge/1.0"})
                    data = json.loads(resp.text)
                    series_list = data.get("series", {}).get("docs", [])
                    if series_list:
                        s = series_list[0]
                        periods = s.get("period", [])
                        values = s.get("value", [])
                        # Data is ascending — take last 6 entries
                        history = []
                        start_idx = max(0, len(periods) - 6)
                        for i in range(len(periods) - 1, start_idx - 1, -1):
                            if i < len(values) and values[i] is not None:
                                history.append({"period": periods[i], "rate": values[i]})
                        if history:
                            results[cc] = {
                                "current_rate": history[0]["rate"],
                                "as_of": history[0]["period"],
                                "history": history,
                            }
            country_names = {"HU": "MNB", "XM": "ECB", "US": "Fed", "CZ": "CNB", "PL": "NBP",
                            "RO": "BNR", "GB": "BoE", "JP": "BoJ", "CH": "SNB", "TR": "TCMB"}
            summary = [f"{country_names.get(cc, cc)}: {r['current_rate']}% ({r['as_of']})"
                       for cc, r in results.items()]
            return json.dumps({"source": "BIS/DBnomics", "summary": summary, "rates": results}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": f"DBnomics API hiba: {e}"})

    if name == "econ_calendar":
        days = min(args.get("days_ahead", 7), 30)
        events = []
        try:
            if FRED_API_KEY:
                from datetime import timedelta
                now_dt = datetime.now(timezone.utc)
                start = now_dt.strftime("%Y-%m-%d")
                end = (now_dt + timedelta(days=days)).strftime("%Y-%m-%d")
                url = (
                    f"https://api.stlouisfed.org/fred/releases/dates"
                    f"?api_key={FRED_API_KEY}&file_type=json"
                    f"&realtime_start={start}&realtime_end={end}"
                    f"&include_release_dates_with_no_data=true&limit=50"
                )
                async with httpx.AsyncClient(timeout=15) as client:
                    resp = await client.get(url)
                data = json.loads(resp.text)
                for rd in data.get("release_dates", []):
                    events.append({
                        "date": rd.get("date", ""),
                        "release": rd.get("release_name", ""),
                        "release_id": rd.get("release_id", ""),
                        "source": "FRED",
                    })
            return json.dumps({"days_ahead": days, "events": events[:30]}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": f"Economic calendar hiba: {e}"})

    return json.dumps({"error": f"Ismeretlen tool: {name}"})


# ── Text-based tool call parser ────────────────────────────────────


def _parse_text_tool_call(content: str):
    """
    Parse text-based tool calls from DeepSeek DSML or Kimi markers.
    Returns (tool_name, args_dict) or None.
    """
    import re

    # DeepSeek DSML format:
    # <｜DSML｜invoke name="calendar_create">
    # <｜DSML｜parameter name="summary" string="true">Value</｜DSML｜parameter>
    dsml_match = re.search(r'invoke name="([^"]+)"', content)
    if dsml_match:
        tool_name = dsml_match.group(1)
        params = {}
        for m in re.finditer(r'parameter name="([^"]+)"[^>]*>([^<]*)<', content):
            params[m.group(1)] = m.group(2)
        if params:
            return tool_name, params

    # Kimi format:
    # <|tool_call_begin|>functions.web_search:2<|tool_call_argument_begin|>{"query": "..."}
    kimi_match = re.search(r'functions\.(\w+).*?argument_begin\|>(\{[^}]+\})', content, re.DOTALL)
    if kimi_match:
        tool_name = kimi_match.group(1)
        try:
            params = json.loads(kimi_match.group(2))
            return tool_name, params
        except json.JSONDecodeError:
            pass

    # Generic: try to find any {"key": "value"} JSON in the text
    json_match = re.search(r'\{[^{}]*"query"[^{}]*\}', content)
    if json_match:
        try:
            params = json.loads(json_match.group())
            return "web_search", params
        except json.JSONDecodeError:
            pass

    return None


# ── Helper functions ───────────────────────────────────────────────


def _fetch_open_tasks(ctx) -> str:
    """Fetch open tasks from Bridge DB for system prompt context."""
    try:
        conn = ctx.get_db()
        rows = conn.execute(
            "SELECT id, title, status, priority FROM tasks "
            "WHERE status IN ('pending', 'in_progress') "
            "ORDER BY created_at DESC LIMIT 10"
        ).fetchall()
        conn.close()
        if not rows:
            return ""
        lines = ["# NYITOTT FELADATOK"]
        for r in rows:
            lines.append(f"- #{r['id']} {r['title']} ({r['status']}, {r['priority']})")
        return "\n".join(lines)
    except Exception:
        return ""


def _store_bridge_reply(ctx, user_text: str, reply: str, agent_id: str):
    """Store Feldwebel reply in Bridge messages DB."""
    try:
        conn = ctx.get_db()
        row = conn.execute(
            "SELECT id FROM messages WHERE sender = 'kommandant' "
            "ORDER BY id DESC LIMIT 1"
        ).fetchone()
        msg_id = row["id"] if row else None
        conn.execute(
            "INSERT INTO messages (timestamp, sender, recipient, subject, message, priority, thread_id, reply_to) "
            "VALUES (?, ?, 'kommandant', ?, ?, 'normal', ?, ?)",
            (datetime.now(timezone.utc).isoformat(),
             f"feldwebel-{agent_id}",
             "Re: Feldwebel válasz",
             reply,
             msg_id, msg_id)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.warning("Failed to store Bridge reply: %s", e)
